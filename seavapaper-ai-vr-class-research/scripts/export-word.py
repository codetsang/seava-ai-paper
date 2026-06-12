#!/usr/bin/env python3
"""Export the结题验收报告 to Word (.docx). Sync content with Report.jsx / reportData.js."""

from __future__ import annotations

import io
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.font_manager import FontProperties, fontManager

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
ASSETS = ROOT / "src" / "assets"

# --- data (sync with src/content/reportData.js) ---
SECTION_WORD_COUNTS = {"part1": 2458, "part2": 660}

SCHEDULE_ROWS = [
    ["2025年5月—7月", "可行性测试与分析，完成测试环境搭建部署", "已完成"],
    ["2025年8月—2026年3月", "正式环境搭建，开展教学试运行", "已完成"],
    ["2026年4月—6月", "教学评价与验收材料整理", "已完成"],
]

ASSESSMENT_ROWS = [
    ["AI系统测评", "30%", "流利度、节奏、发音、内容适合度等"],
    ["教师评分", "30%", "现场表现与口语表达质量"],
    ["学生自评", "20%", "对自身训练过程的反思评价"],
    ["学生互评", "20%", "组间互评与反馈"],
]

EVALUATION_SCORES = {
    "sessions": ["第一次", "第二次", "第三次"],
    "engagement": [70, 76, 81],
    "expression": [64, 72, 78],
    "interest": [67, 74, 80],
}

TRAINING_ROWS = [
    ["第一次口语训练", "部分学生进入偏慢", "以完成基本任务为主", "参与积极性一般"],
    ["第二次口语训练", "多数学生能跟上节奏", "愿意主动开口", "兴趣有所提高"],
    ["第三次口语训练", "多数学生能完成展示", "表达相对完整", "课堂互动更主动"],
]

COMPARE_ROWS = [
    ["角色呈现", "主要靠现场装扮和固定布景", "按任务由AIGC生成虚拟角色形象"],
    ["课前准备", "布景、道具准备耗时较多", "以系统检查、任务说明为主"],
    ["课堂组织", "换装、换景常打断训练", "角色切换在系统内完成"],
    ["训练回看", "录像分散保存，回看不便", "训练录像纳入平台统一回看"],
    ["教辅事务", "布景维护工作较多", "以设备和参数维护为主"],
]


def pick_font() -> str:
    for name in ("Songti SC", "SimSun", "STSong", "PingFang SC", "Microsoft YaHei"):
        if name in {f.name for f in fontManager.ttflist}:
            return name
    return "DejaVu Sans"


def set_run_font(run, name: str, size_pt: float, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def style_document(doc: Document, body_font: str):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)


def add_title(doc: Document, text: str, body_font: str, heading_font: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, heading_font, 16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("结题验收报告")
    set_run_font(run, body_font, 11, color=RGBColor(0x55, 0x55, 0x55))
    sub.paragraph_format.space_after = Pt(18)


def add_h2(doc: Document, text: str, word_count: int | None, heading_font: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, heading_font, 14, bold=True)
    if word_count is not None:
        run2 = p.add_run(f"（{word_count}字）")
        set_run_font(run2, heading_font, 10, color=RGBColor(0x66, 0x66, 0x66))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_h3(doc: Document, text: str, heading_font: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, heading_font, 12, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, body_font: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, body_font, 12)
    p.paragraph_format.first_line_indent = Cm(0.74)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str, body_font: str):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run_font(run, body_font, 10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_run_font(run, body_font, 10.5)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, body_font, 10.5, color=RGBColor(0x44, 0x44, 0x44))
    cap.paragraph_format.space_after = Pt(10)


def add_figure(doc: Document, image_path: Path, caption: str, body_font: str, width_inches: float = 5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, body_font, 10.5, color=RGBColor(0x44, 0x44, 0x44))
    cap.paragraph_format.space_after = Pt(10)


def render_chart_png() -> Path:
    font = pick_font()
    fp = FontProperties(family=font, size=10)
    sessions = EVALUATION_SCORES["sessions"]
    engagement = EVALUATION_SCORES["engagement"]
    expression = EVALUATION_SCORES["expression"]
    interest = EVALUATION_SCORES["interest"]

    x = range(len(sessions))
    width = 0.22
    fig, ax = plt.subplots(figsize=(6.2, 3.2), dpi=150)
    ax.bar([i - width for i in x], engagement, width, label="投入程度", color="#2E75B6")
    ax.bar(x, expression, width, label="表达欲望", color="#548235")
    ax.bar([i + width for i in x], interest, width, label="参与兴趣度", color="#C55A11")
    ax.set_xticks(list(x))
    ax.set_xticklabels(sessions, fontproperties=fp)
    ax.set_ylabel("班级均值", fontproperties=fp)
    ax.set_ylim(55, 90)
    ax.set_yticks(range(55, 91, 5))
    ax.legend(prop=fp, loc="upper left")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bars in ax.containers:
        ax.bar_label(bars, fontsize=8, padding=2)
    fig.tight_layout()
    out = OUT_DIR / "_chart-fig6.png"
    fig.savefig(out, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def build_document() -> Document:
    body_font = pick_font()
    heading_font = "SimHei" if "SimHei" in {f.name for f in fontManager.ttflist} else body_font
    doc = Document()
    style_document(doc, body_font)

    add_title(doc, "基于AIGC数字人的外语情景实验教学项目", body_font, heading_font)
    add_h2(doc, "一、项目结题正文", SECTION_WORD_COUNTS["part1"], heading_font)

    add_h3(doc, "（一）项目概况", heading_font)
    add_para(
        doc,
        '本项目在外语实验室"基于虚拟演播室的情景教学实验项目"基础上实施改造。'
        "该实验通过对表演者视频数据的同步捕捉和渲染开展情景式教学，"
        "在口语训练中发挥了积极作用，但应用场景比较有限。",
        body_font,
    )
    add_para(
        doc,
        "原先实验的问题主要有两方面。一是学生普遍没有受过专业表演训练，演绎不同角色时比较拘束，"
        "不易进入角色，外文戏剧类任务尤为明显；二是受场地条件限制，课堂上服装更换、场景布置、"
        "道具准备等工作耗时较多，情景展示的丰富度和训练深度都受到制约，"
        "多角色、多身份的口语训练往往做不深。",
        body_font,
    )
    add_para(
        doc,
        "项目建设目标是在现有实验条件上搭建基于AIGC引擎的数字人实时生成环境，"
        "对表演者形象进行实时渲染和虚拟角色同步生成，实现多场景、多模态的AIGC情景实验教学，"
        "着重解决依托人工智能技术提升学生口语表达能力和口语综合技能的问题。"
        "项目组在原有绿幕演播条件上完成软硬件升级，接入开源AIGC生成框架，"
        "使课堂能按教学需要生成虚拟角色形象，并据此优化实验流程和评价方式。",
        body_font,
    )

    add_h3(doc, "（二）建设内容与完成情况", heading_font)
    add_para(
        doc,
        "项目建设主要包括软硬件基础条件建设、AIGC引擎及平台部署、情景系统适配和实验教学方案调整。"
        "硬件方面，在原有绿幕环境和摄像采集设备基础上，增配补光及视频处理设备，"
        "改善视频抠像和AIGC数字人生成效果；搭建单工作站运行环境，满足AIGC实时生成需要；"
        "对虚拟演播的环境线路、采集位置和显示终端进行了整理，"
        "绿幕采集、图形处理和显示输出链路经过多轮调试。",
        body_font,
    )
    add_figure(
        doc,
        ASSETS / "photos" / "photo-lab-overview.png",
        "图1 外语实验室AIGC情景实验环境",
        body_font,
    )
    add_figure(
        doc,
        ASSETS / "photos" / "photo-chroma-key-session.png",
        "图2 绿幕抠像与虚拟情景合成现场",
        body_font,
        width_inches=2.75,
    )
    add_para(
        doc,
        "软件方面，采用开源AIGC生成框架作底层能力，在原有情景虚拟仿真实验教学系统上完成适配，"
        "包括音视频流处理、效果器调试、HDMI和DP信号适配、视频采集卡适配等，"
        "将AIGC生成引擎及配套应用接入虚拟情景演播系统，贯通采集、生成、展示和回看各环节。"
        "项目前期已在测试环境中验证数字形象生成的可行性，正式环境部署后用于日常教学。",
        body_font,
    )
    add_figure(
        doc,
        ASSETS / "photos" / "photo-aigc-engine-ui.png",
        "图3 AIGC数字人生成引擎操作界面",
        body_font,
        width_inches=2.75,
    )
    add_para(
        doc,
        '围绕AIGC情景教学和AIGC戏剧类任务，'
        '实行"课前准备+课中实验+课后自评+教师综合评测+AI评测"的教学组织方式。'
        "课前由教辅老师协助学生熟悉系统环境和操作规则，教师布置课程主题和训练任务；"
        "课中学生在AIGC实验环境下完成情景对话、戏剧片段和综合展示等口语训练；"
        "课后学生通过资源平台回看录像，开展自评和互评；"
        "教师结合课堂观察和互评结果进行综合讲评，AI系统对口语流利度、节奏、发音等给出辅助测评。"
        "实验成绩由AI系统测评、教师评分、学生自评和学生互评四部分构成，比例见表4。",
        body_font,
    )
    add_table(doc, ["成绩项目", "占比", "说明"], ASSESSMENT_ROWS, "表4 实验成绩构成", body_font)
    add_para(
        doc,
        "建设过程中，项目组先后完成测试环境部署、正式环境搭建、情景模板导入和教辅培训，"
        "并根据首轮上课情况调整了部分界面操作和默认参数，减少课堂上临时处理的环节。"
        "实验室技术人员参与环境部署和日常维护，按学期维护系统参数和情景模板，"
        "整理操作说明，保障实验在后续学期继续开设。",
        body_font,
    )

    add_h3(doc, "（三）实验教学组织", heading_font)
    add_para(
        doc,
        "升级后的实验按课前准备、课中训练、课后复盘、教师讲评四个主要环节组织实施。"
        "课前，教师说明情景任务和角色要求，布置AIGC戏剧或情景对话主题，"
        "教辅老师协助学生熟悉系统操作、完成参数检查，明确当次练习安排；"
        "课中，教师讲解AIGC环境下的站位、肢体语言和发言规范，"
        "学生在虚拟角色形象支持下进入AIGC情景实验环境，完成口语表达训练；"
        "课后，学生在资源平台回看训练录像，填写自评表并完成组内互评；"
        "讲评时，教师结合课堂观察、录像回放和AI测评报告指出问题，给出成绩。",
        body_font,
    )
    add_para(
        doc,
        "各环节分工与改造前基本一致，主要变化在课中环节和评价方式。"
        "角色由现场装扮优化为AIGC生成，进一步激发了学生的学习和表达热情。",
        body_font,
    )
    add_figure(
        doc,
        ASSETS / "figures" / "fig5-teaching-flow.png",
        "图4 实验教学四环节流程",
        body_font,
    )

    add_h3(doc, "（四）实验教学开展与效果分析", heading_font)
    add_para(
        doc,
        "项目按计划分三个阶段实施：（1）完成可行性测试和测试环境部署；"
        "（2）完成正式环境搭建，并开展多轮教学试运行；"
        "（3）完成教学评价和验收材料整理。各阶段完成情况见表3。",
        body_font,
    )
    add_table(doc, ["时间阶段", "主要工作", "完成情况"], SCHEDULE_ROWS, "表3 项目进度计划完成情况", body_font)
    add_para(
        doc,
        "改造完成后，AIGC情景实验教学在外语实验室按计划开展，组织方式见图5。"
        "本学期实验课安排三次口语训练：第一次熟悉系统和基础情景对话，"
        "第二次加入AIGC戏剧片段练习，第三次做综合展示和讲评。"
        "与改造前相比，课前不必频繁更换布景，同一节课能切换的情景更多，"
        "课堂时间更多用于口语练习；训练录像及时归档，便于课后复盘。"
        "三次训练的任务难度递进。",
        body_font,
    )
    add_figure(
        doc,
        ASSETS / "figures" / "fig4-experiment-flow.png",
        "图5 实验教学流程",
        body_font,
    )
    add_para(
        doc,
        "每次训练结束后，任课教师与教辅老师依据课堂观察表，对班级整体表现按百分制打分，"
        "侧重投入程度、表达欲望和参与兴趣度三项，再取班级均值。"
        "三次训练的统计结果见图6，文字摘要见表1；改造前后组织层面对比见表2。",
        body_font,
    )
    chart_path = render_chart_png()
    add_figure(doc, chart_path, "图6 三次口语训练课堂观察均值", body_font, width_inches=5.2)
    add_table(
        doc,
        ["训练轮次", "投入程度", "表达欲望", "参与兴趣度"],
        TRAINING_ROWS,
        "表1 三次口语训练课堂观察摘要",
        body_font,
    )
    add_table(doc, ["对比项", "改造前", "改造后"], COMPARE_ROWS, "表2 实验项目改造前后主要变化", body_font)
    add_para(
        doc,
        "从图6和表1看，第一次训练三项均值都在70分上下，表达欲望相对偏低，"
        "部分学生对系统和虚拟形象还不习惯，进入角色偏慢；"
        "第二次训练投入程度和参与兴趣度升至74—76分，愿意开口的学生增多；"
        "第三次训练三项均值升至78—81分，多数学生能完成展示任务，课堂互动更主动一些。"
        "各轮次表达欲望均略低于另外两项，与课堂上仍有学生比较紧张、表达不够完整的情况相符。",
        body_font,
    )
    add_para(
        doc,
        "表2反映的主要是组织层面的变化。教师反映，改造后课前准备事务减少，"
        "课堂中换景对训练的打断少了；学生提到较多的是角色切换更方便，便于课后对照练习。"
        "需要说明的是，生成效果会受机房环境、任务设置和学生操作熟练度影响，"
        "个别课时仍有生成延迟或形象与情景不够贴合的情况，还需结合课堂反馈继续调整参数和模板。",
        body_font,
    )

    add_h3(doc, "（五）项目完成情况", heading_font)
    add_para(
        doc,
        "项目计划内容基本完成。AIGC生成所需软硬件已到位并用于日常教学；"
        "实验流程、评价方式和教辅分工均完成调整，形成可复用的AIGC情景实验教学条件一套。",
        body_font,
    )
    add_para(
        doc,
        "后续将继续补充情景模板、完善教辅操作说明，把AIGC戏剧类任务用到更多口语实验内容中，"
        "并结合使用情况优化生成参数，减少课堂等待和反复调试，"
        "让AIGC数字技术更好地服务于外语口语的情景化训练。",
        body_font,
    )

    add_h2(doc, "二、项目特色与创新点及研究成果应用", SECTION_WORD_COUNTS["part2"], heading_font)
    add_h3(doc, "（一）项目特色与创新点", heading_font)
    add_para(
        doc,
        "本项目在原虚拟情景演播系统上完成升级，原有设备、课程内容进一步得到优化，"
        "项目建设投入相对可控，迁移和复用成本较低。相较于动作捕捉类方案，AIGC引擎对场地和硬件的要求更低，"
        "更适合在现有语言实验室条件下逐步改造。",
        body_font,
    )
    add_para(
        doc,
        "AIGC生成能力接入后，课堂能按教学任务实时生成不同的虚拟角色形象，"
        "AIGC情景教学和AIGC戏剧类任务可在同一平台开展，实现多场景、多模态的口语训练。",
        body_font,
    )
    add_para(
        doc,
        '此外，本项目在教学上更注重"人机协同"。教师负责情景设计、过程指导和综合讲评，'
        "AIGC系统承担虚拟角色生成，学生通过自主训练和互评提高参与度。"
        "评价上采用教师评分、学生自评互评与AI系统测评相结合的方式，结果直接用于课堂改进，"
        "没有另设脱离教学实际的评价环节，为口语实验教学改革提供了一条可操作的路径。",
        body_font,
    )

    add_h3(doc, "（二）研究成果应用及效益", heading_font)
    add_para(
        doc,
        "通过此次实验项目的建设，进一步优化了实验室现有项目资源。"
        "AIGC情景环境有助于增加学生的学习兴趣，学生任务完成得比改造前完整一些，"
        "主动开口的意愿有所提高；也有学生把训练内容用到课程展示和英语演讲中。"
        "基础较弱的学生通过录像回看和AI测评反馈，能对照自己的发音和表达问题反复练习。",
        body_font,
    )
    add_para(
        doc,
        "教师方面，口语指导和讲评时间更充裕；"
        "实验教辅工作由布景维护为主转向设备联调和参数维护，日常维护负担减轻。",
        body_font,
    )
    add_para(
        doc,
        "从实验室运行看，原有场地和采集设备利用率提高，布景道具的采购和损耗减少，"
        "有限的人力更多用在教学辅导和设备维护上，现有人员和技术条件能够支撑项目持续开设。"
        "项目实现了复用现有演播条件和开源AIGC框架，"
        "而不是单独建设一套新的数字人系统，对其他口语实验课程和省内高校同类实验教学的"
        "升级改造具有一定的参考意义。项目开放共享实施后，有助于更多学生在外语实验教学中"
        "接触和应用人工智能技术，增强口语学习兴趣和语言自信。",
        body_font,
    )

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / "结题验收报告.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Saved → {out_path}")


if __name__ == "__main__":
    main()
